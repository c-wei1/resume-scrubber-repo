import io
import re
import zipfile
import tempfile
from pathlib import Path

from lxml import etree

from address_identifier import redact_addresses, is_address_line, _REGION_RE
from html_to_docx import parse_quill_html

# City, STATE pattern — matches "San Francisco, CA" style locations
_CITY_STATE_RE = re.compile(
    r'[A-Z][a-zA-Z .\-]{1,30},\s*(' + _REGION_RE.pattern[2:-2] + r')\b'
)


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/package/2006/relationships",
}

IMAGE_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)
HYPERLINK_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


# ═════════════════════════════════════════════════════════════════════
# PII patterns
# ═════════════════════════════════════════════════════════════════════

# Phone: comprehensive first-pass extractor (intl prefixes, parens, ext.)
# Aggressive on purpose — the digit-count validator below rejects false positives.
_PHONE_RE = re.compile(
    r'(?<![\w@.])'                          # not inside emails/decimals/IDs
    r'(?:(?:\+|00)\s?\d{1,3}[\s.\-]?)?'     # optional +XX / 00XX intl prefix
    r'(?:\(\s?\d{1,4}\s?\)[\s.\-]?)?'       # optional (area code)
    r'\d{1,4}(?:[\s.\-]?\d{1,4}){1,5}'      # 2–6 digit groups (≥2 required)
    r'(?:\s?(?:ext|x|extension)\.?\s?\d{1,5})?'  # optional extension
    r'(?![\w])',                            # not followed by a word char
    re.IGNORECASE,
)


def is_plausible_phone(match: str) -> bool:
    """Digit-count gate: real phone numbers have 7–15 digits (E.164 caps at 15)."""
    digits = re.sub(r'\D', '', match)
    return 7 <= len(digits) <= 15


def _redact_phones_sub(text: str) -> str:
    """Regex + validation for plain-string substitution (metadata scrubbing)."""
    return _PHONE_RE.sub(
        lambda m: "" if is_plausible_phone(m.group(0)) else m.group(0),
        text,
    )

_EMAIL_RE = re.compile(
    r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
)

_URL_RE = re.compile(
    r'(?:https?://|www\.|[a-zA-Z0-9\-]+\.(?:com|org|net|io|dev|me|co|info|biz)(?:/|\b))[^\s<>()\[\]"\']+',
    re.IGNORECASE,
)


# ═════════════════════════════════════════════════════════════════════
# Image stripping
# ═════════════════════════════════════════════════════════════════════

def _strip_images(tree):
    """
    Remove image content while preserving any text that may live inside
    drawings/shapes/text boxes.
    """
    root = tree.getroot()

    image_containers = (
        root.xpath("//w:drawing", namespaces=NS)
        + root.xpath("//w:pict", namespaces=NS)
    )

    for el in image_containers:
        has_text = bool(el.xpath(".//*[local-name()='t']"))
        has_textbox = bool(el.xpath(".//*[local-name()='txbx']"))

        if has_text or has_textbox:
            _strip_image_children(el)
        else:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    return tree


def _strip_image_children(el):
    """Remove image-specific children, leaving text-carrying content intact."""
    image_local_names = {
        "pic", "blip", "blipFill", "imagedata", "image",
    }
    for node in list(el.iter()):
        localname = etree.QName(node).localname
        if localname in image_local_names:
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)


# ═════════════════════════════════════════════════════════════════════
# PII redaction
# ═════════════════════════════════════════════════════════════════════
def _find_pii_spans(text: str) -> list[tuple[int, int]]:
    """Return sorted, merged character spans of all PII matches in text."""
    spans: list[tuple[int, int]] = []

    # Email + URL: match directly (unchanged)
    for pattern in (_EMAIL_RE, _URL_RE):
        for m in pattern.finditer(text):
            spans.append(m.span())

    # Phone: regex first-pass, then validate digit count before accepting
    for m in _PHONE_RE.finditer(text):
        if is_plausible_phone(m.group(0)):
            spans.append(m.span())

    # City, STATE locations (e.g. "San Francisco, CA")
    for m in _CITY_STATE_RE.finditer(text):
        spans.append(m.span())

    if not spans:
        return spans
    spans.sort()
    merged = [spans[0]]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
        else:
            merged.append((s, e))
    return merged


def _build_run_map(t_nodes):
    """Map character offsets in combined paragraph text back to <w:t> elements."""
    run_map: list[tuple[int, int, object]] = []
    offset = 0
    for t_el in t_nodes:
        txt = t_el.text or ''
        run_map.append((offset, offset + len(txt), t_el))
        offset += len(txt)
    combined = ''.join((t.text or '') for t in t_nodes)
    return combined, run_map


REDACTED = ""


def _redact_spans_in_runs(spans, run_map, para):
    """Replace PII spans with empty text."""
    if not spans:
        return

    buffers = {}
    for idx, (r_start, r_end, t_el) in enumerate(run_map):
        buffers[idx] = list(t_el.text or '')

    for pii_start, pii_end in spans:
        first_run_seen = False

        for idx, (r_start, r_end, t_el) in enumerate(run_map):
            overlap_start = max(pii_start, r_start)
            overlap_end = min(pii_end, r_end)

            if overlap_start >= overlap_end:
                continue

            buf = buffers[idx]
            local_start = overlap_start - r_start
            local_end = overlap_end - r_start

            for i in range(local_start, local_end):
                buf[i] = ''

            if not first_run_seen:
                buf[local_start] = REDACTED
                first_run_seen = True

    for idx, (r_start, r_end, t_el) in enumerate(run_map):
        t_el.text = ''.join(buffers[idx])


def _strip_pii_for_address_check(text: str) -> str:
    """Remove phone/email/URL patterns from text so their digits
    don't inflate the address scorer, but the rest of the line is
    preserved for an accurate address-vs-not-address decision."""
    out = text
    # Remove phones (validated)
    out = _PHONE_RE.sub(
        lambda m: "" if is_plausible_phone(m.group(0)) else m.group(0),
        out,
    )
    out = _EMAIL_RE.sub("", out)
    out = _URL_RE.sub("", out)
    return out


def _scrub_paragraphs(tree):
    """Paragraph-level PII detection and redaction."""
    root = tree.getroot()
    for para in root.xpath("//w:p", namespaces=NS):
        t_nodes = para.xpath(".//w:t", namespaces=NS)
        if not t_nodes:
            continue

        combined, run_map = _build_run_map(t_nodes)
        if not combined.strip():
            continue

        # Check for address using text with PII patterns stripped out,
        # so phone digits don't inflate the address score.
        cleaned_for_check = _strip_pii_for_address_check(combined)
        if cleaned_for_check.strip() and is_address_line(cleaned_for_check):
            for run in list(para.xpath(".//w:r", namespaces=NS)):
                parent = run.getparent()
                if parent is not None:
                    parent.remove(run)
            continue

        # Redact PII (phone, email, URL) spans
        spans = _find_pii_spans(combined)
        _redact_spans_in_runs(spans, run_map, para)

    _clean_floating_pipes(tree)
    return tree


def _clean_floating_pipes(tree):
    """Remove '|' characters with no meaningful text on both sides."""
    root = tree.getroot()
    for para in root.xpath("//w:p", namespaces=NS):
        t_nodes = para.xpath(".//w:t", namespaces=NS)
        if not t_nodes:
            continue

        combined = ''.join((t.text or '') for t in t_nodes)
        if '|' not in combined:
            continue

        parts = [p.strip() for p in combined.split('|')]
        cleaned = ' | '.join(p for p in parts if p)

        if cleaned == combined:
            continue

        if t_nodes:
            t_nodes[0].text = cleaned
            for t_el in t_nodes[1:]:
                t_el.text = ''


def _scrub_metadata_xml(xml_file: Path):
    tree = etree.parse(str(xml_file))
    root = tree.getroot()

    for tag in root.xpath(".//*"):
        if not tag.text:
            continue

        localname = etree.QName(tag).localname

        if localname in {
            "creator", "lastModifiedBy", "description",
            "subject", "title", "keywords", "category",
        }:
            text = tag.text
            text = _EMAIL_RE.sub("", text)
            text = _redact_phones_sub(text)
            text = _URL_RE.sub("", text)
            tag.text = text.strip()

    tree.write(str(xml_file), encoding="UTF-8", xml_declaration=True)


def _unwrap_hyperlinks(tree):
    """Remove <w:hyperlink> wrappers and replace their visible text."""
    root = tree.getroot()

    for hl in root.xpath("//w:hyperlink", namespaces=NS):
        parent = hl.getparent()
        if parent is None:
            continue

        t_nodes = hl.xpath(".//w:t", namespaces=NS)
        for i, t_el in enumerate(t_nodes):
            t_el.text = REDACTED if i == 0 else ""

        insert_idx = parent.index(hl)
        for child in list(hl):
            hl.remove(child)
            parent.insert(insert_idx, child)
            insert_idx += 1

        parent.remove(hl)

    return tree


def _process_tree(tree):
    _unwrap_hyperlinks(tree)
    _scrub_paragraphs(tree)
    _strip_images(tree)
    return tree


def _strip_rels(rels_file: Path):
    """Remove image and hyperlink relationships from a .rels file."""
    if not rels_file.exists():
        return
    rels_tree = etree.parse(str(rels_file))
    rels_root = rels_tree.getroot()
    remove_types = {IMAGE_REL_TYPE, HYPERLINK_REL_TYPE}
    for rel in list(rels_root):
        if rel.get("Type") in remove_types:
            rels_root.remove(rel)
    rels_tree.write(str(rels_file), encoding="UTF-8", xml_declaration=True)


# ═════════════════════════════════════════════════════════════════════
# Scrubbing pipeline
# ═════════════════════════════════════════════════════════════════════

def process_docx(input_bytes: bytes) -> io.BytesIO:
    with tempfile.TemporaryDirectory() as _tmpdir:
        tmpdir = Path(_tmpdir)

        with zipfile.ZipFile(io.BytesIO(input_bytes), "r") as z:
            z.extractall(tmpdir)

        word_dir = tmpdir / "word"

        # Scrub metadata
        docprops_dir = tmpdir / "docProps"
        for xml_file in docprops_dir.glob("*.xml"):
            print("Scrubbing metadata:", xml_file)
            _scrub_metadata_xml(xml_file)

        # Main document
        doc_xml = word_dir / "document.xml"
        if doc_xml.exists():
            tree = etree.parse(str(doc_xml))
            _process_tree(tree)
            tree.write(str(doc_xml), encoding="UTF-8", xml_declaration=True)

        # Headers & footers
        for xml_file in (
            list(word_dir.glob("header*.xml"))
            + list(word_dir.glob("footer*.xml"))
        ):
            tree = etree.parse(str(xml_file))
            _process_tree(tree)
            tree.write(str(xml_file), encoding="UTF-8", xml_declaration=True)

        # Relationships
        _strip_rels(word_dir / "_rels" / "document.xml.rels")
        for header_rels in (word_dir / "_rels").glob("header*.xml.rels"):
            _strip_rels(header_rels)
        for footer_rels in (word_dir / "_rels").glob("footer*.xml.rels"):
            _strip_rels(footer_rels)

        # Media files
        media_dir = word_dir / "media"
        if media_dir.exists():
            for f in media_dir.iterdir():
                if f.is_file():
                    f.unlink()

        # Repack
        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as z:
            for f in tmpdir.rglob("*"):
                if f.is_file():
                    z.write(f, f.relative_to(tmpdir))
        output.seek(0)

        return output


def _prepend_user_info(
    docx_bytes: io.BytesIO,
    name: str,
    title: str,
    department: str,
    responsibilities: str = "",
) -> io.BytesIO:
    """
    DEPRECATED: Use _insert_experience_entry instead.
    Prepend name, title, department, and responsibilities as paragraphs at
    the top of the first page in 12pt Times New Roman.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        with zipfile.ZipFile(docx_bytes, "r") as z:
            z.extractall(tmpdir)

        doc_xml = tmpdir / "word" / "document.xml"
        tree = etree.parse(str(doc_xml))
        root = tree.getroot()

        body = root.find(f"{{{W}}}body")
        if body is None:
            docx_bytes.seek(0)
            return docx_bytes

        # Build label-value pairs for Name, Job Title, Department
        info_lines = []
        if name:
            info_lines.append(("Name:", name))
        if title:
            info_lines.append(("Job Title:", title))
        if department:
            info_lines.append(("Department:", department))

        def _make_run(parent, text, bold=False, italic=False, underline=False):
            """Add a <w:r> with 12pt Times New Roman and optional formatting."""
            run = etree.SubElement(parent, f"{{{W}}}r")
            rPr = etree.SubElement(run, f"{{{W}}}rPr")
            rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
            rFonts.set(f"{{{W}}}ascii", "Times New Roman")
            rFonts.set(f"{{{W}}}hAnsi", "Times New Roman")
            sz = etree.SubElement(rPr, f"{{{W}}}sz")
            sz.set(f"{{{W}}}val", "24")
            szCs = etree.SubElement(rPr, f"{{{W}}}szCs")
            szCs.set(f"{{{W}}}val", "24")
            if bold:
                etree.SubElement(rPr, f"{{{W}}}b")
            if italic:
                etree.SubElement(rPr, f"{{{W}}}i")
            if underline:
                u = etree.SubElement(rPr, f"{{{W}}}u")
                u.set(f"{{{W}}}val", "single")
            t = etree.SubElement(run, f"{{{W}}}t")
            t.text = text
            t.set(f"{{{W}}}space", "preserve")
            return run

        def _make_para(text, bold=False):
            """Create a <w:p> with 12pt Times New Roman, no extra spacing."""
            para = etree.Element(f"{{{W}}}p")
            pPr = etree.SubElement(para, f"{{{W}}}pPr")
            spacing = etree.SubElement(pPr, f"{{{W}}}spacing")
            spacing.set(f"{{{W}}}after", "0")
            spacing.set(f"{{{W}}}line", "240")
            spacing.set(f"{{{W}}}lineRule", "auto")
            _make_run(para, text, bold=bold)
            return para

        def _make_rich_para(runs_data, list_type=None):
            """Create a <w:p> with multiple formatted runs."""
            para = etree.Element(f"{{{W}}}p")
            pPr = etree.SubElement(para, f"{{{W}}}pPr")
            spacing = etree.SubElement(pPr, f"{{{W}}}spacing")
            spacing.set(f"{{{W}}}after", "0")
            spacing.set(f"{{{W}}}line", "240")
            spacing.set(f"{{{W}}}lineRule", "auto")

            # Add bullet/number prefix
            if list_type == 'bullet':
                combined = ''.join(r['text'] for r in runs_data).strip()
                if not combined.startswith('\u2022'):
                    _make_run(para, '\u2022 ')
            elif list_type == 'ordered':
                # The counter prefix is added by the caller
                pass

            for run_data in runs_data:
                text = run_data['text']
                if not text:
                    continue
                _make_run(
                    para, text,
                    bold=run_data.get('bold', False),
                    italic=run_data.get('italic', False),
                    underline=run_data.get('underline', False),
                )
            return para

        # Insert in reverse order so final order is correct
        if responsibilities:
            parsed = parse_quill_html(responsibilities)
            if parsed:
                ol_counter = 0
                rich_paras = []
                for p in parsed:
                    if p['list_type'] == 'ordered':
                        ol_counter += 1
                        # Prepend number to the first run
                        first_run = p['runs'][0] if p['runs'] else None
                        if first_run:
                            prefix_text = first_run['text'].strip()
                            if not prefix_text.startswith(f'{ol_counter}.'):
                                p['runs'].insert(0, {
                                    'text': f'{ol_counter}. ',
                                    'bold': False, 'italic': False, 'underline': False
                                })
                    else:
                        if p['list_type'] != 'ordered':
                            ol_counter = 0
                    rich_paras.append(p)

                for p in reversed(rich_paras):
                    body.insert(0, _make_rich_para(p['runs'], p['list_type']))
            else:
                # Fallback: plain text
                resp_lines = [l for l in responsibilities.splitlines() if l.strip()]
                for line in reversed(resp_lines):
                    bullet_line = line if line.strip().startswith("\u2022") else f"\u2022 {line.strip()}"
                    body.insert(0, _make_para(bullet_line))
            body.insert(0, _make_para("Current Responsibilities at Gilead:", bold=True))

        for label, value in reversed(info_lines):
            body.insert(0, _make_para(value))
            body.insert(0, _make_para(label, bold=True))

        tree.write(
            str(doc_xml),
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

        output = io.BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as z:
            for f in tmpdir.rglob("*"):
                if f.is_file():
                    z.write(f, f.relative_to(tmpdir))
        output.seek(0)
        return output


# ── Experience‐header keywords (must stay in sync with model_section_parser) ──
_EXP_ALIASES = {
    "experience", "work experience", "professional experience",
    "employment", "employment history", "work history", "career history",
    "professional background", "relevant experience",
}


def _is_experience_header(text: str) -> bool:
    """Return True if `text` looks like an experience section header."""
    raw = (text or "").strip().strip(":").strip()
    if not raw or len(raw) > 45:
        return False
    low = raw.lower()
    for alias in _EXP_ALIASES:
        if low == alias or low.startswith(alias + " ") or low.startswith(alias + ":"):
            return True
        letters_only = re.sub(r"[^a-z]", "", low)
        alias_letters = alias.replace(" ", "")
        if alias in low and len(letters_only) <= len(alias_letters) + 6:
            return True
    return False


def _insert_experience_entry(
    docx_bytes: io.BytesIO,
    title: str,
    department: str,
    start_date: str,
    responsibilities: str = "",
) -> io.BytesIO:
    """
    Insert a Gilead experience entry at the beginning of the experience
    section in the resume.

    Strategy (3-step fallback):
      1. Find an experience section header → insert right after it
      2. Use NER model to detect experience section → insert before first entry
      3. Last resort → insert at top of document

    Line 1 (bold): "Title, Department, StartDate-Present"
    Lines 2+: Current responsibilities (from Quill HTML)
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # Save to a temp file so python-docx and TextExtractor can work with it.
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        tmp_path = tmpdir / "source.docx"
        docx_bytes.seek(0)
        tmp_path.write_bytes(docx_bytes.read())

        # ── 1. Build paragraph specs from user input ─────────────
        left_parts = [p for p in [title, department, "Gilead Sciences Inc."] if p]
        left_text = ",  ".join(left_parts) if left_parts else ""
        date_text = f"{start_date} - Present" if start_date else ""

        para_specs = []

        # "Current Responsibilities" heading first (bold + underlined)
        para_specs.append({"runs": [{"text": "Current Responsibilities",
                                      "bold": True, "italic": False,
                                      "underline": True}],
                            "list_type": None, "date_text": ""})

        # Job Title, Department, Gilead with date right-aligned
        if left_text or date_text:
            para_specs.append({"runs": [{"text": left_text, "bold": True,
                                          "italic": False, "underline": False}],
                                "date_text": date_text,
                                "list_type": None})

        if responsibilities:
            parsed = parse_quill_html(responsibilities)
            if parsed:
                ol_counter = 0
                for p in parsed:
                    if p["list_type"] == "ordered":
                        ol_counter += 1
                        p["runs"].insert(0, {"text": f"{ol_counter}. ",
                                             "bold": False, "italic": False,
                                             "underline": False})
                    else:
                        ol_counter = 0
                    para_specs.append(p)
            else:
                for line in responsibilities.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    if not line.startswith("\u2022"):
                        line = f"\u2022 {line}"
                    para_specs.append({"runs": [{"text": line, "bold": False,
                                                  "italic": False, "underline": False}],
                                        "list_type": None})

        if not para_specs:
            docx_bytes.seek(0)
            return docx_bytes

        # ── 2. Open ONE Document instance for both detection & insertion ─
        doc = Document(str(tmp_path))
        body = doc.element.body

        # Iterate ALL <w:p> elements (including those nested in tables/SDTs)
        # — this matches what TextExtractor does for the template path.
        all_paragraphs = list(body.iter(qn("w:p")))

        # Helper: walk up from any element to find its body-level ancestor
        # (the direct child of <w:body> that contains it).
        def _body_level_ancestor(el):
            current = el
            while current is not None:
                parent = current.getparent()
                if parent is body:
                    return current
                current = parent
            return None

        insert_before_el = None
        # Track the parent where we'll insert (could be body or a table cell)
        insert_parent = None

        # Step 1: Scan all paragraphs for an experience header
        for p_el in all_paragraphs:
            text = "".join(t.text or "" for t in p_el.iter(qn("w:t")))
            if _is_experience_header(text):
                # Insert right after the header (at the start of the section).
                insert_before_el = p_el.getnext()
                insert_parent = p_el.getparent()
                break

        # Step 2: Fallback — use NER ModelSectionParser
        if insert_before_el is None and insert_parent is None:
            try:
                from parser_get_text import TextExtractor
                from model_section_parser import ModelSectionParser

                pairs = TextExtractor.extract_pairs(tmp_path)
                _model_path = str(Path(__file__).resolve().parent / "resume_ner_model")
                parser = ModelSectionParser(_model_path, use_model=True)
                sections = parser.find_sections(pairs)

                exp_pairs = sections.get("experience", [])
                if exp_pairs:
                    # Get the text of the first experience paragraph
                    first_exp_text = exp_pairs[0][0] if isinstance(exp_pairs[0], (tuple, list)) else ""
                    first_exp_text = first_exp_text.strip()

                    if first_exp_text:
                        # Match by text content against our doc's paragraphs
                        for p_el in all_paragraphs:
                            p_text = "".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip()
                            if p_text == first_exp_text:
                                # Insert before this paragraph, within its parent
                                insert_before_el = p_el
                                insert_parent = p_el.getparent()
                                break
            except Exception as e:
                print(f"[_insert_experience_entry] NER fallback failed: {e}")

        # Step 3: Last resort — insert at top of body
        if insert_before_el is None and insert_parent is None:
            children = list(body)
            insert_before_el = children[0] if children else None
            insert_parent = body

        # ── 3. Detect font and size from the experience section ──
        from collections import Counter
        from docx.shared import Inches
        from docx.oxml import OxmlElement

        # Match experience section paragraphs in python-docx to get
        # resolved font properties (with style inheritance).
        exp_font_counter = Counter()
        exp_size_counter = Counter()
        in_exp = False
        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not in_exp:
                if _is_experience_header(p_text):
                    in_exp = True
                    continue
            else:
                if p_text and len(p_text) <= 45 and not _is_experience_header(p_text):
                    from parser_get_section import SectionParser
                    canon = SectionParser._canonical(p_text)
                    if canon and canon in SectionParser._get_canonical_lookup():
                        break
                for run in p.runs:
                    text = (run.text or "").strip()
                    if not text:
                        continue
                    char_count = len(text)
                    if run.font.name:
                        exp_font_counter[run.font.name] += char_count
                    if run.font.size:
                        exp_size_counter[run.font.size] += char_count

        # Pick font: experience section → style default → Times New Roman
        if exp_font_counter:
            primary_font = exp_font_counter.most_common(1)[0][0]
        else:
            try:
                primary_font = doc.styles["Normal"].font.name or "Times New Roman"
            except Exception:
                primary_font = "Times New Roman"

        # Pick size: experience section (capped at 12pt) → style default (capped) → 12pt
        MAX_BODY_SIZE = Pt(12)
        if exp_size_counter:
            primary_size = exp_size_counter.most_common(1)[0][0]
            if primary_size > MAX_BODY_SIZE:
                primary_size = MAX_BODY_SIZE
        else:
            try:
                default_size = doc.styles["Normal"].font.size
                if default_size and default_size <= MAX_BODY_SIZE:
                    primary_size = default_size
                else:
                    primary_size = MAX_BODY_SIZE
            except Exception:
                primary_size = MAX_BODY_SIZE

        # ── 4. Create and insert paragraphs ──────────────────────
        for spec in para_specs:
            new_para = doc.add_paragraph()
            new_para.paragraph_format.space_after = Pt(0)
            new_para.paragraph_format.space_before = Pt(0)

            is_bullet = spec.get("list_type") == "bullet"

            if is_bullet:
                new_para.paragraph_format.left_indent = Inches(0.25)
                combined = "".join(r["text"] for r in spec.get("runs", [])).strip()
                if not combined.startswith("\u2022"):
                    run = new_para.add_run("\u2022 ")
                    run.font.name = primary_font
                    run.font.size = primary_size

            for run_data in spec.get("runs", []):
                text = run_data.get("text", "")
                if not text:
                    continue
                run = new_para.add_run(text)
                run.font.name = primary_font
                run.font.size = primary_size
                run.bold = run_data.get("bold", False)
                run.italic = run_data.get("italic", False)
                run.underline = run_data.get("underline", False)

            # Right-align date text via a tab stop at the right margin
            date_text = spec.get("date_text", "")
            if date_text:
                # Add a right-aligned tab stop to paragraph properties
                pPr = new_para._element.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    new_para._element.insert(0, pPr)
                tabs = OxmlElement("w:tabs")
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), "right")
                tab.set(qn("w:pos"), "9360")  # 6.5 inches in twips (standard margin)
                tab.set(qn("w:leader"), "none")
                tabs.append(tab)
                pPr.append(tabs)

                # Add a tab character run, then the date run
                tab_run = new_para.add_run("\t")
                tab_run.font.name = primary_font
                tab_run.font.size = primary_size
                date_run = new_para.add_run(date_text)
                date_run.font.name = primary_font
                date_run.font.size = primary_size
                date_run.bold = True

            # Move from end of body (where add_paragraph puts it) to target
            body.remove(new_para._element)
            if insert_before_el is not None:
                insert_before_el.addprevious(new_para._element)
            elif insert_parent is not None:
                insert_parent.append(new_para._element)
            else:
                body.append(new_para._element)

        # Add a blank line after the inserted entry
        blank_para = doc.add_paragraph()
        blank_para.paragraph_format.space_after = Pt(0)
        blank_para.paragraph_format.space_before = Pt(0)
        body.remove(blank_para._element)
        if insert_before_el is not None:
            insert_before_el.addprevious(blank_para._element)
        elif insert_parent is not None:
            insert_parent.append(blank_para._element)
        else:
            body.append(blank_para._element)

        # ── 4. Save and return ───────────────────────────────────
        out_path = tmpdir / "output.docx"
        doc.save(str(out_path))

        output = io.BytesIO(out_path.read_bytes())
        output.seek(0)
        return output
