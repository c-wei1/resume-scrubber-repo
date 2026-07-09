from pathlib import Path


try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(
        "Error: python-docx not installed. "
        "Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)



NEWLINE_TAG = "[NEWLINE]"


class TextExtractor:
    """Extract plain text from DOCX files."""

    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

    _BLOCKED_ANCESTORS = {
        # qn("w:drawing"),
        # qn("w:pict"),
        # qn("w:txbxContent"),
        f"{{{MC_NS}}}Fallback",
    }


    _T_TAG = qn("w:t")
    _BR_TAG = qn("w:br")

    @classmethod
    def _is_blocked(cls, element) -> bool:
        parent = element.getparent()
        while parent is not None:
            if parent.tag in cls._BLOCKED_ANCESTORS:
                return True
            parent = parent.getparent()
        return False

    @classmethod
    def _paragraph_text(cls, paragraph) -> str:
        """
        Concatenate all <w:t> text under this paragraph, including
        SDT/content-controls. Preserve <w:br/> line breaks as inline
        [NEWLINE] markers. Skip text inside drawings/text boxes.
        """
        parts = []

        for el in paragraph._element.iter(cls._T_TAG, cls._BR_TAG):
            if cls._is_blocked(el):
                continue

            if el.tag == cls._T_TAG:
                if el.text:
                    parts.append(el.text)
            # else:
            #     # <w:br/>: inline line break inside a paragraph
            #     parts.append(f" {NEWLINE_TAG} ")

        return "".join(parts).strip()

    @classmethod
    def _extract_paragraph_line(cls, paragraph) -> str:
        text = cls._paragraph_text(paragraph)
        return text # if text else NEWLINE_TAG

    @classmethod
    def _extract_table(cls, table) -> list:
        lines = []
        seen_tc = set()

        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if id(tc) in seen_tc:
                    continue
                seen_tc.add(id(tc))

                for paragraph in cell.paragraphs:
                    lines.append(cls._extract_paragraph_line(paragraph))

                for nested in cell.tables:
                    lines.extend(cls._extract_table(nested))

        return lines

    @classmethod
    def _paragraph_text_from_element(cls, p_element) -> str:
        parts = []
        for el in p_element.iter(cls._T_TAG, cls._BR_TAG):
            if cls._is_blocked(el):
                continue
            if el.tag == cls._T_TAG:
                if el.text:
                    parts.append(el.text)
            # else:  # w:br
            #     parts.append(f" {NEWLINE_TAG} ")
        return "".join(parts).strip()

    @classmethod
    def _iter_body_paragraphs(cls, doc):
        """
        Yield every <w:p> in document order, including paragraphs inside
        tables, text boxes, drawings, and content controls.
        """
        body = doc.element.body
        for p in body.iter(qn("w:p")):
            yield p


    @classmethod
    def extract_docx(cls, docx_path: Path) -> str:
        doc = Document(str(docx_path))

        lines = []
        for p in cls._iter_body_paragraphs(doc):
            text = cls._paragraph_text_from_element(p)
            lines.append(text) # if text else NEWLINE_TAG)
        return "\n".join(lines)

    @classmethod
    def extract(cls, file_path: Path) -> str:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return cls.extract_docx(file_path)

        if suffix == ".txt":
            text = file_path.read_text(encoding="utf-8")
            return "\n".join(text)

        raise ValueError(f"Unsupported file type: {suffix}")
