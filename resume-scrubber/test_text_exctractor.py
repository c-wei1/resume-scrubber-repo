from docx import Document

doc = Document("/Users/cwei1/Library/CloudStorage/OneDrive-GileadSciences/Desktop/cv_checker_script/resumes/L.Wilkins Resume 25-Jul-2024.docx")
for para in doc.paragraphs:
    print(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)