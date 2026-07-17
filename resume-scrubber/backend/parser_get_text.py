# from pathlib import Path


# try:
#     from docx import Document
#     from docx.oxml.ns import qn
# except ImportError:
#     print(
#         "Error: python-docx not installed. "
#         "Install with: pip install python-docx",
#         file=sys.stderr,
#     )
#     sys.exit(1)



# NEWLINE_TAG = "[NEWLINE]"


# class TextExtractor:
#     """Extract plain text from DOCX files."""

#     MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

#     _BLOCKED_ANCESTORS = {
#         # qn("w:drawing"),
#         # qn("w:pict"),
#         # qn("w:txbxContent"),
#         f"{{{MC_NS}}}Fallback",
#     }


#     _T_TAG = qn("w:t")
#     _BR_TAG = qn("w:br")

#     @classmethod
#     def _is_blocked(cls, element) -> bool:
#         parent = element.getparent()
#         while parent is not None:
#             if parent.tag in cls._BLOCKED_ANCESTORS:
#                 return True
#             parent = parent.getparent()
#         return False

#     @classmethod
#     def _paragraph_text(cls, paragraph) -> str:
#         """
#         Concatenate all <w:t> text under this paragraph, including
#         SDT/content-controls. Preserve <w:br/> line breaks as inline
#         [NEWLINE] markers. Skip text inside drawings/text boxes.
#         """
#         parts = []

#         for el in paragraph._element.iter(cls._T_TAG, cls._BR_TAG):
#             if cls._is_blocked(el):
#                 continue

#             if el.tag == cls._T_TAG:
#                 if el.text:
#                     parts.append(el.text)
#             # else:
#             #     # <w:br/>: inline line break inside a paragraph
#             #     parts.append(f" {NEWLINE_TAG} ")

#         return "".join(parts).strip()

#     @classmethod
#     def _extract_paragraph_line(cls, paragraph) -> str:
#         text = cls._paragraph_text(paragraph)
#         return text # if text else NEWLINE_TAG

#     @classmethod
#     def _extract_table(cls, table) -> list:
#         lines = []
#         seen_tc = set()

#         for row in table.rows:
#             for cell in row.cells:
#                 tc = cell._tc
#                 if id(tc) in seen_tc:
#                     continue
#                 seen_tc.add(id(tc))

#                 for paragraph in cell.paragraphs:
#                     lines.append(cls._extract_paragraph_line(paragraph))

#                 for nested in cell.tables:
#                     lines.extend(cls._extract_table(nested))

#         return lines

#     @classmethod
#     def _paragraph_text_from_element(cls, p_element) -> str:
#         parts = []
#         for el in p_element.iter(cls._T_TAG, cls._BR_TAG):
#             if cls._is_blocked(el):
#                 continue
#             if el.tag == cls._T_TAG:
#                 if el.text:
#                     parts.append(el.text)
#             # else:  # w:br
#             #     parts.append(f" {NEWLINE_TAG} ")
#         return "".join(parts).strip()

#     @classmethod
#     def _iter_body_paragraphs(cls, doc):
#         """
#         Yield every <w:p> in document order, including paragraphs inside
#         tables, text boxes, drawings, and content controls.
#         """
#         body = doc.element.body
#         for p in body.iter(qn("w:p")):
#             yield p


#     @classmethod
#     def _clean_line(cls, text: str) -> str:
#         """Replace control/whitespace unicode characters with their visual equivalents."""
#         text = text.replace("\t", "    ")       # tab → spaces
#         text = text.replace("\xa0", " ")        # non-breaking space → space
#         text = text.replace("\u200b", "")       # zero-width space → remove
#         text = text.replace("\u00ad", "")       # soft hyphen → remove
#         text = text.replace("\u2011", "-")      # non-breaking hyphen → hyphen
#         text = text.replace("\u2013", "–")      # en dash
#         text = text.replace("\u2014", "—")      # em dash
#         text = text.replace("\u2018", "'")      # left single quote
#         text = text.replace("\u2019", "'")      # right single quote
#         text = text.replace("\u201c", '"')      # left double quote
#         text = text.replace("\u201d", '"')      # right double quote
#         text = text.replace("\u2026", "...")     # ellipsis
#         text = text.replace("\u2022", "•")      # bullet (keep as-is, already visible)
#         return text

#     @classmethod
#     def _get_paragraph_text(cls, p_el):
#         """
#         Get text from a <w:p> element, but only from <w:t> nodes that
#         belong directly to this paragraph — not from nested paragraphs
#         inside text boxes or drawings.
#         """
#         P_TAG = qn("w:p")
#         parts = []
#         for t in p_el.iter(cls._T_TAG):
#             # Check if this <w:t> is inside a NESTED <w:p> (text box, etc.)
#             # by walking up from the <w:t> to see if we hit another <w:p>
#             # before reaching our p_el
#             parent = t.getparent()
#             nested = False
#             while parent is not None and parent is not p_el:
#                 if parent.tag == P_TAG:
#                     nested = True
#                     break
#                 parent = parent.getparent()
#             if not nested and t.text:
#                 parts.append(t.text)
#         return "".join(parts)

#     @classmethod
#     def extract_docx(cls, docx_path: Path) -> str:
#         doc = Document(str(docx_path))

#         lines = []
#         last_text_key = None
#         body = doc.element.body

#         # Iterate ALL <w:p> elements in document order — this includes
#         # paragraphs inside tables, SDTs (content controls), text boxes, etc.
#         for p_el in body.iter(qn("w:p")):
#             text = cls._get_paragraph_text(p_el)
#             cleaned = cls._clean_line(text)

#             # Deduplicate consecutive identical lines (merged table cells,
#             # repeated text-box content)
#             text_key = cleaned.strip()
#             if text_key and text_key == last_text_key:
#                 continue
#             if text_key:
#                 last_text_key = text_key
#             else:
#                 last_text_key = None

#             lines.append(cleaned)

#         return "\n".join(lines)

#     @classmethod
#     def extract(cls, file_path: Path) -> str:
#         file_path = Path(file_path)
#         suffix = file_path.suffix.lower()

#         if suffix == ".docx":
#             return cls.extract_docx(file_path)

#         if suffix == ".txt":
#             text = file_path.read_text(encoding="utf-8")
#             return "\n".join(text)

#         raise ValueError(f"Unsupported file type: {suffix}")
"""
TextExtractor
=============

Single source of truth for reading paragraph text out of a DOCX.

Two entry points:

    extract_pairs(path)   yields (cleaned_text, <w:p> element) tuples
                          in document order. Use this for anything that
                          needs to align text-level decisions with the
                          underlying XML (e.g., section detection that
                          later feeds template population).

    extract(path)         returns a single '\n'-joined string. Kept for
                          backward compatibility with existing callers.
"""

import sys
from pathlib import Path
from typing import Iterable, Iterator, List, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print(
        "Error: python-docx not installed. "
        "Install with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


NEWLINE_TAG = "[NEWLINE]"


class TextExtractor:
    """Extract plain text from DOCX files as (text, xml) pairs."""

    MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

    _BLOCKED_ANCESTORS = {
        f"{{{MC_NS}}}Fallback",
    }

    _T_TAG = qn("w:t")
    _BR_TAG = qn("w:br")
    _P_TAG = qn("w:p")

    # ─────────────────────────────────────────────────────────────
    # Low-level helpers
    # ─────────────────────────────────────────────────────────────
    @classmethod
    def _is_blocked(cls, element) -> bool:
        parent = element.getparent()
        while parent is not None:
            if parent.tag in cls._BLOCKED_ANCESTORS:
                return True
            parent = parent.getparent()
        return False

    @classmethod
    def _get_paragraph_text(cls, p_el) -> str:
        """
        Concatenate every <w:t> that belongs directly to this <w:p>.
        Skips <w:t> nodes that live inside a *nested* <w:p> (text boxes,
        drawings that contain their own paragraphs) so the caller sees
        exactly one text string per top-level paragraph.
        """
        parts: List[str] = []
        for t in p_el.iter(cls._T_TAG):
            if cls._is_blocked(t):
                continue

            # Is this <w:t> inside a nested <w:p>?
            nested = False
            parent = t.getparent()
            while parent is not None and parent is not p_el:
                if parent.tag == cls._P_TAG:
                    nested = True
                    break
                parent = parent.getparent()

            if not nested and t.text:
                parts.append(t.text)
        return "".join(parts)

    @classmethod
    def _clean_line(cls, text: str) -> str:
        """Normalize control / whitespace unicode into their visual forms."""
        text = text.replace("\t", "    ")
        text = text.replace("\xa0", " ")
        text = text.replace("\u200b", "")
        text = text.replace("\u00ad", "")
        text = text.replace("\u2011", "-")
        text = text.replace("\u2013", "–")
        text = text.replace("\u2014", "—")
        text = text.replace("\u2018", "'")
        text = text.replace("\u2019", "'")
        text = text.replace("\u201c", '"')
        text = text.replace("\u201d", '"')
        text = text.replace("\u2026", "...")
        text = text.replace("\u2022", "•")
        return text

    # ─────────────────────────────────────────────────────────────
    # Primary API — (text, xml) pairs
    # ─────────────────────────────────────────────────────────────
    @classmethod
    def extract_pairs(cls, docx_path: Path) -> List[Tuple[str, "etree._Element"]]:
        """
        Yield one (cleaned_text, <w:p> element) tuple per paragraph in
        document order, including paragraphs inside tables and SDTs.

        Consecutive duplicate lines (merged table cells, repeated text-box
        content) are collapsed to a single pair. The XML element for the
        collapsed line is the FIRST occurrence — this matches what a
        reader would visually see.
        """
        doc = Document(str(docx_path))
        pairs: List[Tuple[str, "etree._Element"]] = []
        last_key: str = ""

        body = doc.element.body
        for p_el in body.iter(qn("w:p")):
            raw = cls._get_paragraph_text(p_el)
            cleaned = cls._clean_line(raw)
            key = cleaned.strip()

            # Deduplicate consecutive identical lines
            if key and key == last_key:
                continue

            last_key = key if key else ""
            pairs.append((cleaned, p_el))

        return pairs

    # ─────────────────────────────────────────────────────────────
    # Convenience — text-only view
    # ─────────────────────────────────────────────────────────────
    @classmethod
    def extract_docx(cls, docx_path: Path) -> str:
        return "\n".join(text for text, _ in cls.extract_pairs(docx_path))

    @classmethod
    def extract(cls, file_path: Path) -> str:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            return cls.extract_docx(file_path)
        if suffix == ".txt":
            return file_path.read_text(encoding="utf-8")
        raise ValueError(f"Unsupported file type: {suffix}")
